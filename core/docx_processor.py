import docx
from typing import List, Dict, Any, Union
from io import BytesIO
from docx.document import Document
from docx.table import _Cell
from .models import DocumentPatch

class DocxProcessor:
    """
    负责处理DOCX文件的核心类。
    功能包括：解析文档结构、为元素生成ID、提取内容与格式、以及应用修改。
    所有操作现在都在内存中进行。
    """

    @staticmethod
    def extract_structure_with_ids(file_stream: BytesIO) -> Dict[str, Any]:
        """
        从内存中的文件流解析DOCX文件，提取其结构，并为每个可编辑元素生成唯一ID。

        :param file_stream: 包含.docx文件内容的BytesIO流。
        :return: 一个代表文档结构的字典。
        """
        document = docx.Document(file_stream)
        structure = {"elements": []}
        element_counter = 0

        # 遍历文档的所有顶级元素（段落和表格）
        for element in document.element.body:
            if element.tag.endswith('p'):  # 如果是段落
                p = docx.text.paragraph.Paragraph(element, document)
                paragraph_id = f"p_{element_counter}"
                structure["elements"].append({
                    "id": paragraph_id,
                    "type": "paragraph",
                    "text": p.text,
                    "style": p.style.name,
                })
                element_counter += 1
            elif element.tag.endswith('tbl'):  # 如果是表格
                table = docx.table.Table(element, document)
                table_id = f"tbl_{element_counter}"
                table_data = {
                    "id": table_id,
                    "type": "table",
                    "rows": []
                }
                for i, row in enumerate(table.rows):
                    row_data = {"cells": []}
                    for j, cell in enumerate(row.cells):
                        cell_id = f"{table_id}_r{i}c{j}"
                        row_data["cells"].append({
                            "id": cell_id,
                            "text": cell.text,
                        })
                    table_data["rows"].append(row_data)
                structure["elements"].append(table_data)
                element_counter += 1
        
        return structure

    @staticmethod
    def apply_patches(original_stream: BytesIO, new_stream: BytesIO, patches: List[DocumentPatch]):
        """
        将一系列修改（补丁）应用到内存中的原始DOCX文件流，并将结果写入新的流。

        :param original_stream: 包含原始.docx文件内容的BytesIO流。
        :param new_stream: 用于写入修改后文件内容的BytesIO流。
        :param patches: 一个包含修改指令的列表。
        """
        document = docx.Document(original_stream)
        patches_dict = {p.element_id: p.new_content for p in patches}

        element_counter = 0
        
        # 再次遍历文档元素以应用修改
        for element in document.element.body:
            if element.tag.endswith('p'):
                current_id = f"p_{element_counter}"
                if current_id in patches_dict:
                    p = docx.text.paragraph.Paragraph(element, document)
                    DocxProcessor._replace_paragraph_text(p, patches_dict[current_id])
                element_counter += 1
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, document)
                table_id = f"tbl_{element_counter}"
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell_id = f"{table_id}_r{i}c{j}"
                        if cell_id in patches_dict:
                            # 清空单元格并填充新内容
                            cell.text = str(patches_dict[cell_id])
                element_counter += 1

        document.save(new_stream)

    @staticmethod
    def _replace_paragraph_text(p: docx.text.paragraph.Paragraph, new_text: str):
        """
        替换段落的文本，同时尽量保留原始格式。
        
        该方法采用以下策略来保留格式：
        1. 如果段落只有一个run，直接替换文本内容，完全保留格式
        2. 如果段落有多个runs，尝试保留格式分布
        3. 如果新文本较短，使用第一个run的格式
        4. 如果新文本较长，按比例分配格式

        :param p: 要修改的段落对象。
        :param new_text: 新的文本内容。
        """
        if not p.runs:
            # 如果段落没有runs，创建一个新的run
            p.add_run(new_text)
            return
        
        # 保存原始runs的格式信息
        original_runs_info = []
        original_text = ""
        
        for run in p.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': None,
                'style': None
            }
            
            # 保存字体颜色（如果有）
            if run.font.color.rgb:
                run_info['font_color'] = run.font.color.rgb
            
            original_runs_info.append(run_info)
            original_text += run.text
        
        # 清空现有的runs
        for run in p.runs[:]:
            r = run._r
            r.getparent().remove(r)
        
        # 策略1: 如果原段落只有一个run，直接替换并保留所有格式
        if len(original_runs_info) == 1:
            new_run = p.add_run(new_text)
            DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 策略2: 如果新文本为空，不添加任何内容
        if not new_text.strip():
            return
        
        # 策略3: 如果新文本很短或原文本很短，使用第一个run的格式
        if len(new_text) <= 50 or len(original_text) <= 10:
            new_run = p.add_run(new_text)
            DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 策略4: 尝试按比例分配格式
        DocxProcessor._distribute_formatting_proportionally(p, new_text, original_runs_info, original_text)

    @staticmethod
    def _apply_run_formatting(run, format_info: Dict[str, Any]):
        """
        将格式信息应用到run对象上。
        
        :param run: docx run对象
        :param format_info: 包含格式信息的字典
        """
        try:
            if format_info.get('bold') is not None:
                run.bold = format_info['bold']
            if format_info.get('italic') is not None:
                run.italic = format_info['italic']
            if format_info.get('underline') is not None:
                run.underline = format_info['underline']
            if format_info.get('font_name'):
                run.font.name = format_info['font_name']
            if format_info.get('font_size'):
                run.font.size = format_info['font_size']
            if format_info.get('font_color'):
                run.font.color.rgb = format_info['font_color']
        except Exception:
            # 如果应用格式失败，忽略错误继续执行
            pass

    @staticmethod
    def _distribute_formatting_proportionally(p, new_text: str, original_runs_info: List[Dict], original_text: str):
        """
        按比例分配格式到新文本中。
        
        :param p: 段落对象
        :param new_text: 新文本
        :param original_runs_info: 原始runs的格式信息
        :param original_text: 原始文本
        """
        if not original_text:
            # 如果原文本为空，使用第一个run的格式
            new_run = p.add_run(new_text)
            if original_runs_info:
                DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 计算每个原始run在文本中的位置比例
        current_pos = 0
        run_positions = []
        
        for run_info in original_runs_info:
            run_text = run_info['text']
            start_ratio = current_pos / len(original_text) if len(original_text) > 0 else 0
            end_ratio = (current_pos + len(run_text)) / len(original_text) if len(original_text) > 0 else 1
            
            run_positions.append({
                'start_ratio': start_ratio,
                'end_ratio': end_ratio,
                'format_info': run_info
            })
            current_pos += len(run_text)
        
        # 根据比例将新文本分段并应用相应格式
        new_text_len = len(new_text)
        last_end = 0
        
        for i, pos_info in enumerate(run_positions):
            start_pos = int(pos_info['start_ratio'] * new_text_len)
            end_pos = int(pos_info['end_ratio'] * new_text_len)
            
            # 确保不会超出边界
            start_pos = max(start_pos, last_end)
            end_pos = min(end_pos, new_text_len)
            
            # 如果是最后一个run，确保包含所有剩余文本
            if i == len(run_positions) - 1:
                end_pos = new_text_len
            
            if start_pos < end_pos:
                segment_text = new_text[start_pos:end_pos]
                if segment_text:  # 只有当片段不为空时才创建run
                    new_run = p.add_run(segment_text)
                    DocxProcessor._apply_run_formatting(new_run, pos_info['format_info'])
                    last_end = end_pos
        
        # 如果还有剩余文本，使用最后一个run的格式
        if last_end < new_text_len:
            remaining_text = new_text[last_end:]
            if remaining_text:
                new_run = p.add_run(remaining_text)
                if original_runs_info:
                    DocxProcessor._apply_run_formatting(new_run, original_runs_info[-1]) 